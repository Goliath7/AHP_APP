import datetime
import tkinter as tk
from tkinter import ttk, messagebox, scrolledtext
from tkcalendar import DateEntry
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from shared_data import STATIONS, LETTER_SUPERVISORS, output_path

SUPERVISORS = LETTER_SUPERVISORS

DAY_SHIFT = "дневная смена с 8:00 по 17:00"
NIGHT_SHIFT = "ночную смену с 00:00 по 06:00"

class LetterApp:
    def __init__(self, container):
        self.root = container
        if hasattr(self.root, "title"):
            self.root.title("Создание исходящего письма")
        if hasattr(self.root, "geometry"):
            self.root.geometry("820x740")

        self.work_list = []           # список выбранных станций
        self.date_rows = []           # список (frame, DateEntry) для удаления

        self.create_widgets()

    def create_widgets(self):
        frame = ttk.Frame(self.root, padding="12")
        frame.pack(fill=tk.BOTH, expand=True)

        row = 0

        ttk.Label(frame, text="Исх. № (только число):").grid(row=row, column=0, sticky="e", pady=5, padx=5)
        self.num_var = tk.StringVar(value="263")
        ttk.Entry(frame, textvariable=self.num_var, width=12).grid(row=row, column=1, sticky="w")
        row += 1

        ttk.Label(frame, text="Договор:").grid(row=row, column=0, sticky="e", pady=5, padx=5)
        self.contract_var = tk.StringVar(value="№ 4313м от 10.06.2024 г.")
        ttk.Entry(frame, textvariable=self.contract_var, width=45).grid(row=row, column=1, sticky="w")
        row += 1

        ttk.Separator(frame, orient='horizontal').grid(row=row, column=0, columnspan=3, sticky="ew", pady=10)
        row += 1

        ttk.Label(frame, text="Станция:").grid(row=row, column=0, sticky="e", pady=6, padx=5)
        self.station_combo = ttk.Combobox(frame, values=STATIONS, width=38, state="readonly")
        self.station_combo.grid(row=row, column=1, sticky="w")
        self.station_combo.set(STATIONS[0])
        row += 1

        ttk.Label(frame, text="Смена:").grid(row=row, column=0, sticky="e", pady=5, padx=5)
        self.shift_var = tk.StringVar(value="Дневная")
        ttk.Radiobutton(frame, text="Дневная (8:00–17:00)", variable=self.shift_var, value="Дневная").grid(row=row, column=1, sticky="w")
        row += 1
        ttk.Radiobutton(frame, text="Ночная  (00:00–06:00)", variable=self.shift_var, value="Ночная").grid(row=row, column=1, sticky="w")
        row += 1

        ttk.Label(frame, text="Даты:").grid(row=row, column=0, sticky="ne", pady=6, padx=5)

        self.date_frame = ttk.Frame(frame)
        self.date_frame.grid(row=row, column=1, sticky="w", pady=4)

        self.date_rows = []
        self.add_date_row()  # первая строка

        ttk.Button(frame, text="+ ещё дата", command=self.add_date_row).grid(row=row+1, column=1, sticky="w", pady=6)
        row += 2

        ttk.Label(frame, text="Руководитель:").grid(row=row, column=0, sticky="e", pady=6, padx=5)
        self.sup_combo = ttk.Combobox(frame, values=[f"{n} — {p}" for n,p in SUPERVISORS], width=45, state="readonly")
        self.sup_combo.grid(row=row, column=1, sticky="w")
        self.sup_combo.set(self.sup_combo["values"][0])
        row += 1

        ttk.Button(frame, text="Добавить станцию в список", command=self.add_station).grid(row=row, column=1, pady=12, sticky="w")
        row += 1

        ttk.Label(frame, text="Добавленные станции:").grid(row=row, column=0, sticky="ne", pady=6, padx=5)
        self.listbox = scrolledtext.ScrolledText(frame, width=78, height=10, wrap=tk.WORD, font=("Consolas", 10))
        self.listbox.grid(row=row, column=1, pady=4, sticky="nsew")
        row += 1

        ttk.Button(frame, text="Удалить выбранную строку", command=self.delete_selected).grid(row=row, column=1, sticky="w", pady=6)

        frame.columnconfigure(1, weight=1)
        frame.rowconfigure(row-1, weight=1)

        ttk.Button(frame, text="Создать документ (.docx)", command=self.generate_doc).grid(row=row+2, column=1, pady=20, sticky="w")

    def add_date_row(self):
        row_frame = ttk.Frame(self.date_frame)
        row_frame.pack(fill=tk.X, pady=2)

        today = datetime.date.today()
        entry = DateEntry(
            row_frame,
            date_pattern='dd.mm.yy',
            width=12,
            year=today.year, month=today.month, day=today.day
        )
        entry.pack(side=tk.LEFT, padx=(0, 6))

        btn_del = ttk.Button(row_frame, text="−", width=2,
                             command=lambda rf=row_frame: self.remove_date_row(rf))
        btn_del.pack(side=tk.LEFT)

        self.date_rows.append((row_frame, entry))

        # Отключаем кнопку удаления, если осталась только одна строка
        if len(self.date_rows) == 1:
            btn_del.config(state="disabled")
        else:
            # Активируем кнопку на предыдущей строке, если она была отключена
            if len(self.date_rows) > 1:
                self.date_rows[-2][0].winfo_children()[1].config(state="normal")

    def remove_date_row(self, row_frame):
        if len(self.date_rows) <= 1:
            return

        for i, (rf, entry) in enumerate(self.date_rows):
            if rf == row_frame:
                rf.destroy()
                del self.date_rows[i]
                break

        # Если осталась одна строка — отключаем её кнопку удаления
        if len(self.date_rows) == 1:
            self.date_rows[0][0].winfo_children()[1].config(state="disabled")

    def add_station(self):
        station = self.station_combo.get().strip()
        if not station:
            return

        shift = self.shift_var.get()
        is_night = shift == "Ночная"
        shift_text = NIGHT_SHIFT if is_night else DAY_SHIFT

        dates = []
        for _, entry in self.date_rows:
            try:
                d = entry.get_date().strftime("%d.%m")
                dates.append(d)
            except:
                continue

        if not dates:
            messagebox.showwarning("Ошибка", "Выберите хотя бы одну дату")
            return

        if is_night:
            # Для ночных смен: предыдущий день / выбранный день
            night_ranges = []
            for date_str in dates:
                try:
                    day, month = map(int, date_str.split('.'))
                    # Год не важен для вычисления предыдущего дня
                    current = datetime.date(2000, month, day)
                    prev = current - datetime.timedelta(days=1)
                    prev_str = prev.strftime("%d.%m")
                    night_ranges.append(f"{prev_str}/{date_str}")
                except:
                    night_ranges.append(date_str)  # на случай ошибки
            formatted = ", ".join(night_ranges)
            dates_str = f"{formatted}, в ночные смены с 00:00 по 06:00"
        else:
            formatted = ", ".join(dates)
            dates_str = f"{formatted}  {shift_text}"

        sup_text = self.sup_combo.get()
        if "—" not in sup_text:
            messagebox.showwarning("Ошибка", "Выберите руководителя")
            return
        sup_name, sup_phone = [x.strip() for x in sup_text.split("—", 1)]

        line = f"Ст. «{station}»:   {dates_str}\nРуководитель работ: {sup_name}. Тел {sup_phone}."
        self.listbox.insert(tk.END, line + "\n\n")
        self.work_list.append((station, dates_str, sup_name, sup_phone))

    def delete_selected(self):
        try:
            lines = self.listbox.get("1.0", tk.END).splitlines()
            if not lines:
                return

            cursor_line = int(float(self.listbox.index("insert"))) - 1
            cursor_line = max(0, min(cursor_line, len(lines) - 1))

            # Если курсор на пустой строке между блоками, смещаемся к ближайшему блоку вниз.
            while cursor_line < len(lines) and not lines[cursor_line].strip():
                cursor_line += 1
            if cursor_line >= len(lines):
                return

            start = cursor_line
            while start > 0 and lines[start - 1].strip():
                start -= 1

            end = cursor_line
            while end + 1 < len(lines) and lines[end + 1].strip():
                end += 1

            text = "\n".join(lines[start:end + 1]).strip()
            if not text:
                return

            for i, (station, dates_str, name, phone) in enumerate(self.work_list):
                check = f"Ст. «{station}»:   {dates_str}\nРуководитель работ: {name}. Тел {phone}."
                if check.strip() == text:
                    del self.work_list[i]
                    break

            delete_until_line = end + 1
            if end + 1 < len(lines) and not lines[end + 1].strip():
                delete_until_line = end + 2

            self.listbox.delete(f"{start + 1}.0", f"{delete_until_line}.0")
        except Exception:
            pass

    def generate_doc(self):
        if not self.work_list:
            messagebox.showwarning("Ошибка", "Добавьте хотя бы одну станцию")
            return

        try:
            num = int(self.num_var.get().strip())
        except:
            messagebox.showerror("Ошибка", "Исх. № должен быть числом")
            return

        outgoing_number = f"{num}/1"
        filename = f"Исх_{num}_адресат_справа.docx"

        today = datetime.date.today()
        months_ru = {
            1: "января", 2: "февраля", 3: "марта", 4: "апреля", 5: "мая", 6: "июня",
            7: "июля", 8: "августа", 9: "сентября", 10: "октября", 11: "ноября", 12: "декабря"
        }
        date_str = f"{today.day} {months_ru[today.month]} {today.year} г."

        contract = self.contract_var.get().strip() or "№ 4313м от 10.06.2024 г."

        doc = Document()

        section = doc.sections[0]
        section.top_margin    = Cm(1.7)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.8)
        section.right_margin  = Cm(1.5)

        # Логотип (если есть)
        try:
            p = doc.add_paragraph()
            run = p.add_run()
            run.add_picture("image1.png", width=Inches(5.9))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except:
            pass

        doc.add_paragraph()

        p = doc.add_paragraph(f"Исх. № {outgoing_number} от {date_str}")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.runs[0]
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        run.font.size = Pt(11)
        run.bold = True

        p = doc.add_paragraph(
            "Первому заместителю начальника\n"
            "Метрополитена – Начальнику Дирекции инфраструктуры\n"
            "ГУП «Московский метрополитен»\n"
            "Бочанаеву А. А."
        )
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for r in p.runs:
            r.font.name = 'Times New Roman'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            r.font.size = Pt(11)
            r.bold = True

        for _ in range(2):
            doc.add_paragraph()

        p = doc.add_paragraph("Уважаемый Антон Алиевич!")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.name = 'Times New Roman'
        p.runs[0].font.size = Pt(11)

        doc.add_paragraph()

        main_text = (
            f"Во исполнение обязательств по договору {contract}, для выполнения работ "
            "по оформлению архитектурно-художественного освещения прошу выделить "
            "сотрудников Отдела технического надзора службы электроснабжения (ОТН Э) "
            "и Отдела технического надзора службы пассажирских обустройств (ОТН СПО), "
            "а также сотрудников Отдела строительного контроля службы электроснабжения "
            "(СК ОТН Э) для присутствия на объекте:"
        )
        p = doc.add_paragraph(main_text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(6)
        p.runs[0].font.name = 'Times New Roman'
        p.runs[0].font.size = Pt(11)

        doc.add_paragraph()

        for station, dates_str, sup_name, sup_phone in self.work_list:
            p = doc.add_paragraph(f"Ст. «{station}»:   {dates_str}")
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.runs[0].bold = True
            p.runs[0].font.name = 'Times New Roman'
            p.runs[0].font.size = Pt(11)

            p = doc.add_paragraph(f"Руководитель работ: {sup_name}. Тел {sup_phone}.")
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.runs[0].font.name = 'Times New Roman'
            p.runs[0].font.size = Pt(11)

        for _ in range(4):
            doc.add_paragraph()

        p = doc.add_paragraph("Генеральный директор")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.runs[0].font.name = 'Times New Roman'
        p.runs[0].font.size = Pt(11)

        p = doc.add_paragraph("ООО «Азбука Света»")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.runs[0].font.name = 'Times New Roman'
        p.runs[0].font.size = Pt(11)

        p = doc.add_paragraph("Макаров Артём Андреевич")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.runs[0].font.name = 'Times New Roman'
        p.runs[0].font.size = Pt(11)
        p.runs[0].bold = True

        try:
            file_path = output_path(filename)
            doc.save(file_path)
            messagebox.showinfo("Успех", f"Файл сохранён:\n{file_path}")
        except Exception as e:
            messagebox.showerror("Ошибка сохранения", str(e))


if __name__ == "__main__":
    root = tk.Tk()
    app = LetterApp(root)
    root.mainloop()
