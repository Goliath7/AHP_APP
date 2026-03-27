import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt

from shared_data import output_path

DAY_SHIFT = "дневная смена с 8:00 по 17:00"
NIGHT_SHIFT = "ночную смену с 00:00 по 06:00"


def format_dates_for_shift(dates, shift_type):
    if not dates:
        raise ValueError("Добавьте хотя бы одну дату")

    unique_dates = sorted(set(dates))
    formatted_dates = [d.strftime("%d.%m") for d in unique_dates]

    if shift_type == "Ночная":
        night_ranges = []
        for date_obj, date_str in zip(unique_dates, formatted_dates):
            prev = date_obj - datetime.timedelta(days=1)
            night_ranges.append(f"{prev.strftime('%d.%m')}/{date_str}")
        return f"{', '.join(night_ranges)}, в ночные смены с 00:00 по 06:00"

    return f"{', '.join(formatted_dates)}  {DAY_SHIFT}"


def generate_letter_document(num, contract, work_list):
    if not work_list:
        raise ValueError("Добавьте хотя бы одну станцию")

    try:
        num = int(str(num).strip())
    except Exception as exc:
        raise ValueError("Исх. № должен быть числом") from exc

    outgoing_number = f"{num}/1"
    filename = f"Исх_{num}_адресат_справа.docx"
    file_path = output_path(filename)

    today = datetime.date.today()
    months_ru = {
        1: "января",
        2: "февраля",
        3: "марта",
        4: "апреля",
        5: "мая",
        6: "июня",
        7: "июля",
        8: "августа",
        9: "сентября",
        10: "октября",
        11: "ноября",
        12: "декабря",
    }
    date_str = f"{today.day} {months_ru[today.month]} {today.year} г."
    contract = contract.strip() or "№ 4313м от 10.06.2024 г."

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(1.5)

    # Логотип, если есть в проекте.
    logo = Path(__file__).resolve().parent / "image1.png"
    if logo.exists():
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_picture(str(logo), width=Inches(5.9))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    p = doc.add_paragraph(f"Исх. № {outgoing_number} от {date_str}")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.runs[0]
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(11)
    run.bold = True

    p = doc.add_paragraph(
        "Первому заместителю начальника\n"
        "Метрополитена – Начальнику Дирекции инфраструктуры\n"
        "ГУП «Московский метрополитен»\n"
        "Бочанаеву А. А."
    )
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in p.runs:
        r.font.name = "Times New Roman"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        r.font.size = Pt(11)
        r.bold = True

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph("Уважаемый Антон Алиевич!")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.name = "Times New Roman"
    p.runs[0].font.size = Pt(11)

    doc.add_paragraph()

    main_text = (
        f"Во исполнение обязательств по договору {contract}, для выполнения работ "
        "по оформлению архитектурно-художественного освещения прошу выделить "
        "сотрудников Отдела технического надзора службы электроснабжения (ОТН Э) "
        "и Отдела технического надзора службы пассажирских обустройств (ОТН СПО), "
        "а также сотрудников Отдела строительного контроля службы электроснабжения "
        "(СК ОТН Э) для присутствия на объекте:"
    )
    p = doc.add_paragraph(main_text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(6)
    p.runs[0].font.name = "Times New Roman"
    p.runs[0].font.size = Pt(11)

    doc.add_paragraph()

    for station, dates_str, sup_name, sup_phone in work_list:
        p = doc.add_paragraph(f"Ст. «{station}»:   {dates_str}")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.runs[0].bold = True
        p.runs[0].font.name = "Times New Roman"
        p.runs[0].font.size = Pt(11)

        p = doc.add_paragraph(f"Руководитель работ: {sup_name}. Тел {sup_phone}.")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.runs[0].font.name = "Times New Roman"
        p.runs[0].font.size = Pt(11)

    for _ in range(4):
        doc.add_paragraph()

    for line in ("Генеральный директор", "ООО «Азбука Света»", "Макаров Артём Андреевич"):
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.runs[0].font.name = "Times New Roman"
        p.runs[0].font.size = Pt(11)
        if "Макаров" in line:
            p.runs[0].bold = True

    doc.save(file_path)
    return str(file_path)
