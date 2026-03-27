try:
    import tkinter as tk
    from tkinter import ttk, messagebox
    from tkcalendar import DateEntry
except Exception:  # pragma: no cover - for headless/web environments
    tk = None
    ttk = None
    messagebox = None
    DateEntry = None
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from shared_data import STATIONS, LEADERS_FULL, output_path


# ====================== ФУНКЦИИ ДЛЯ WORD ======================
def set_font(run, size=11, bold=False, underline=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.underline = underline
    run.font.name = 'Calibri'
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Calibri')
    rFonts.set(qn('w:hAnsi'), 'Calibri')
    rPr.append(rFonts)


def create_naryad(station, leader_full_declined, leader_short, start_date_str, end_date_str, filename):
    doc = Document()

    # ====================== НАСТРОЙКИ ДОКУМЕНТА ======================
    section = doc.sections[0]
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    normal_style = doc.styles['Normal']
    pf = normal_style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1

    # ====================== ПЕРВАЯ СТРАНИЦА ======================
    # Заголовок
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run('НАРЯД\nна производство работ сторонними организациями\nна объектах ГУП «Московский метрополитен»\n'), 
             size=11, bold=True)
    set_font(p.add_run('(заполняется в двух экземплярах организацией, ведущей производство работ)'), 
             size=9)

    p = doc.add_paragraph('При производстве работ соблюдай\n Правила безопасности и Указания технадзора')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(p.runs[0], size=11)

    doc.add_paragraph()

    # Организация и телефон
    p = doc.add_paragraph()
    set_font(p.add_run('Организация: '), size=11)
    run = p.add_run('ООО «Азбука Света»')
    set_font(run, underline=True)

    p = doc.add_paragraph()
    set_font(p.add_run('Телефон: '), size=11)
    run = p.add_run('8(800) 555-41-44')
    set_font(run, underline=True)

    doc.add_paragraph()

    # Номер наряда
    p = doc.add_paragraph('НАРЯД №  ')
    set_font(p.runs[0], size=11)
    run = p.add_run('_______')
    set_font(run, underline=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Руководитель работ (склонённое ФИО)
    p = doc.add_paragraph('Руководителю работ ')
    run = p.add_run(leader_full_declined)
    set_font(run, underline=True)

    # Основной текст работ
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p.add_run('С бригадой в составе ').font.size = Pt(11)
    run = p.add_run('  7  ')
    set_font(run, underline=True)
    p.add_run(' чел. на основании совместного Приказа от 21.10.2024 № УД-07-2127/24/17 ').font.size = Pt(11)

    run = p.add_run(f'поручается выполнить: работы по прокладке кабельных линий, монтажу несущих конструкций путем сверления, '
                    f'установки щита электропитания и управления освещением, установки анкеров или шпилек на химический анкер, '
                    f'установки металлических кронштейнов и светильников, их коммутации/юстировки для оформления '
                    f'архитектурно-художественного освещения наземного вестибюля станции «{station}» Московского метрополитена')
    set_font(run, underline=True)

    p = doc.add_paragraph('(наименование, место проведения, описание работ)')
    set_font(p.runs[0], size=9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Даты
    p = doc.add_paragraph('Начало работ ')
    run = p.add_run(start_date_str)
    set_font(run, underline=True)

    p = doc.add_paragraph('Окончание работ ')
    run = p.add_run(end_date_str)
    set_font(run, underline=True)

    # Ответственный руководитель
    p = doc.add_paragraph('Ответственный руководитель работ ')
    run = p.add_run('Алимханов Идрис Махмудович')
    set_font(run, underline=True)
    p.add_run('\n+7(915) 237-72-73')

    doc.add_paragraph()

    # Инструктаж
    p = doc.add_paragraph(
        'Весь персонал, включенный в состав бригады, проинструктирован по технике безопасности и знанию '
        'Инструкции «О порядке производства работ сторонними организациями на объектах ГУП «Московский метрополитен», '
        'и по состоянию здоровья допущен к производству работ на высоте на станции.'
    )
    set_font(p.runs[0], size=11)

    doc.add_paragraph()

    # Подписи
    p = doc.add_paragraph('Наряд и инструктаж по данной работе получил руководитель работ')
    p = doc.add_paragraph(f'_______________________________________________ / {leader_short} /')

    p = doc.add_paragraph('Ответственный руководитель работ ________________________________ / Алимханов И. М. /')

    p = doc.add_paragraph('Наряд выдал генеральный директор ____________________________ /Макаров А. А. /')

    p = doc.add_paragraph('м.п.')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # ====================== ВТОРАЯ СТРАНИЦА (ОБОРОТНАЯ) ======================
    doc.add_page_break()

    p = doc.add_paragraph('(оборотная сторона наряда. Заполняется Дистанцией, допускающей к производству работ)')
    set_font(p.runs[0], size=11)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph(f'Работу разрешается проводить с     {start_date_str}                                           по     {end_date_str}')
    set_font(p.runs[0], size=11)

    p = doc.add_paragraph('с        00 час. 00 мин.                  до      06 час. 00 мин.')
    set_font(p.runs[0], size=11)
    p = doc.add_paragraph('с        08 час. 00 мин.                  до      17 час. 00 мин.')
    set_font(p.runs[0], size=11)
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run('при условии:').font.size = Pt(11)

    run = p.add_run(
        ' соблюдения совместного Приказа от 21.10.2024 № УД-07-2127/24/17; '
        '«Инструкций о порядке производства работ сторонними организациями в эксплуатируемых сооружениях Московского метрополитена», введенной приказом от 25.10.2023г. №УД-07-2339/23; '
        '«Инструкции по организации безопасного проведения огневых работ на объектах ГУП «Московский метрополитен», введенной указанием от 29.01.2021 № УД-07-218/21 + изменение от 04.06.2021 № УД-07-2054/21; '
        '«Правил безопасности при строительстве подземных сооружений» ПБ 03-428-02, утвержденных постановлением Мосгортехнадзора 02.11.2001г.; '
        '«Инструкции о пропускном и внутриобъектовом режимах на объекте ГУП «Московский метрополитен»», утверждённой приказом от 12.03.2021 № УД-07-962/21 со всеми изменениями и дополнениями; '
        '«Правил технической эксплуатации метрополитена РФ» и «Правил противопожарного режима в Российской Федерации», утверждённых постановлением Правительства РФ от 16.09.2020 № 1479.'
    )
    set_font(run, size=11, underline=True)

    doc.add_paragraph('Разрешение на производство работ выдал: _____________________________________________')
    doc.add_paragraph('Задание на ведение технадзора получил: ______________________________________________')
    doc.add_paragraph('Задание на ведение технадзора получил: ______________________________________________')
    doc.add_paragraph('Задание на ведение технадзора получил: ______________________________________________')
    doc.add_paragraph()

    p = doc.add_paragraph(
        'Место для согласования с другими службами (право согласования имеют начальники дистанций и их заместители. '
        'При выдаче согласования должна быть указана должность, фамилия, дата)'
    )
    set_font(p.runs[0], size=11)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p = doc.add_paragraph('______________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________')

    p = doc.add_paragraph('Оформление начала и окончания работ')
    set_font(p.runs[0], size=11)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Таблица
    table = doc.add_table(rows=12, cols=10)
    table.style = 'Table Grid'

    cell_start = table.cell(0, 0)
    cell_start.merge(table.cell(0, 4))
    run = cell_start.paragraphs[0].add_run('Начало работ')
    set_font(run, size=11, bold=True)
    cell_start.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    cell_end = table.cell(0, 5)
    cell_end.merge(table.cell(0, 9))
    run = cell_end.paragraphs[0].add_run('Окончание работ')
    set_font(run, size=11, bold=True)
    cell_end.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    headers = ['Дата', 'Часы', 'Минуты', 'Производитель работ (подпись)', 'Технадзор (подпись)',
               'Дата', 'Часы', 'Минуты', 'Производитель работ (подпись)', 'Технадзор (подпись)']

    for i, h in enumerate(headers):
        cell = table.cell(1, i)
        run = cell.paragraphs[0].add_run(h)
        set_font(run, size=10, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for r in range(2, 12):
        for c in range(10):
            table.cell(r, c).text = ''

    file_path = output_path(filename)
    doc.save(file_path)
    return str(file_path)


# ====================== СКЛОНЕНИЕ ФИО ======================
def decline_fio(full_name):
    """Простое склонение ФИО в родительный падеж"""
    parts = full_name.strip().split()
    if len(parts) < 3:
        return full_name

    last, first, patron = parts[0], parts[1], parts[2]

    # Склоняем фамилию
    if last.endswith(('ов', 'ев', 'ин', 'ын')):
        last_decl = last + 'у'
    elif last.endswith(('ский', 'цкий')):
        last_decl = last[:-2] + 'ому'
    elif last.endswith(('ая', 'яя')):
        last_decl = last[:-2] + 'ой'
    elif last.endswith(('ий', 'ой')):
        last_decl = last[:-2] + 'ому'
    elif last.endswith('ь'):
        last_decl = last[:-1] + 'ю'
    elif last.endswith('а'):
        last_decl = last[:-1] + 'е'
    else:
        last_decl = last + 'у'

    # Склоняем имя и отчество
    if first.endswith(('й', 'ь')):
        first_decl = first[:-1] + 'ю'
    elif first.endswith('а'):
        first_decl = first[:-1] + 'е'
    else:
        first_decl = first + 'у'

    if patron.endswith(('ич', 'ыч')):
        patron_decl = patron + 'у'
    elif patron.endswith('на'):
        patron_decl = patron[:-1] + 'е'
    else:
        patron_decl = patron + 'у'

    return f"{last_decl} {first_decl} {patron_decl}"


def format_date_ru(date):
    months = ["января", "февраля", "марта", "апреля", "мая", "июня",
              "июля", "августа", "сентября", "октября", "ноября", "декабря"]
    return f"{date.day} {months[date.month-1]} {date.year} года"


def build_naryad_filename(station, leader_full):
    parts = leader_full.split()
    leader_short = f"{parts[0]} {parts[1][0]}.{parts[2][0]}." if len(parts) >= 3 else leader_full
    station_short = station.split()[0] if station.split() else "Станция"
    filename = f"Наряд_{station_short}_{leader_short.rstrip('.')}.docx"
    return filename, leader_short


def generate_naryad_document(station, leader_full, start_date, end_date):
    station = station.strip()
    leader_full = leader_full.strip()
    if not station:
        raise ValueError("Выберите станцию")
    if not leader_full:
        raise ValueError("Выберите руководителя работ")
    if start_date > end_date:
        raise ValueError("Дата начала не может быть позже даты окончания")

    leader_declined = decline_fio(leader_full)
    filename, leader_short = build_naryad_filename(station, leader_full)
    start_str = format_date_ru(start_date)
    end_str = format_date_ru(end_date)
    return create_naryad(station, leader_declined, leader_short, start_str, end_str, filename)


# ====================== GUI ======================
def update_filename_label(*args):
    station = station_var.get().strip()
    leader_idx = leader_combo.current()
    
    if leader_idx == -1 or not station:
        filename_label.config(text="Файл сохранится как: ...")
        return

    leader_full = leaders[leader_idx]
    filename, _ = build_naryad_filename(station, leader_full)
    filename_label.config(text=f"Файл сохранится как: {filename}")


def generate():
    station = station_var.get().strip()
    if not station:
        messagebox.showerror("Ошибка", "Выберите станцию!")
        return

    leader_idx = leader_combo.current()
    if leader_idx == -1:
        messagebox.showerror("Ошибка", "Выберите руководителя работ!")
        return

    leader_full = leaders[leader_idx]
    start_date = start_cal.get_date()
    end_date = end_cal.get_date()

    try:
        created = generate_naryad_document(station, leader_full, start_date, end_date)
        messagebox.showinfo("Готово!", f"Документ успешно создан!\nФайл: {created}")
    except ValueError as e:
        messagebox.showerror("Ошибка", str(e))
    except Exception as e:
        messagebox.showerror("Ошибка", str(e))


# ====================== ДАННЫЕ ======================
stations = STATIONS
leaders = LEADERS_FULL

# ====================== ОКНО ======================
def run_gui():
    global root, station_var, station_combo, leader_combo, start_cal, end_cal, filename_label
    if tk is None or ttk is None or DateEntry is None or messagebox is None:
        raise RuntimeError("Tkinter GUI недоступен в этой среде")

    root = tk.Tk()
    root.title("Генератор Наряда — Московский метрополитен")
    root.geometry("850x720")
    root.resizable(False, False)

    tk.Label(root, text="Генератор Наряда", font=("Arial", 16, "bold")).pack(pady=10)

    tk.Label(root, text="Станция:", font=("Arial", 10, "bold")).pack(anchor="w", padx=40)
    station_var = tk.StringVar()
    station_combo = ttk.Combobox(root, textvariable=station_var, values=stations, width=80, state="readonly", font=("Arial", 10))
    station_combo.pack(padx=40, pady=5)
    station_combo.set("Курская КЛ")

    tk.Label(root, text="Руководитель работ (кому поручается):", font=("Arial", 10, "bold")).pack(anchor="w", padx=40, pady=(15,5))
    leader_combo = ttk.Combobox(root, values=leaders, width=80, state="readonly", font=("Arial", 10))
    leader_combo.pack(padx=40, pady=5)
    leader_combo.set("Матаев Антон Викторович")

    # Календари
    today = datetime.now().date()
    end_default = today + timedelta(days=10)

    tk.Label(root, text="Начало работ:", font=("Arial", 10, "bold")).pack(anchor="w", padx=40, pady=(20,5))
    start_cal = DateEntry(
        root,
        width=25,
        background="darkblue",
        foreground="white",
        borderwidth=2,
        date_pattern="dd.mm.yyyy",
        year=today.year,
        month=today.month,
        day=today.day,
    )
    start_cal.pack(padx=40, pady=5)

    tk.Label(root, text="Окончание работ:", font=("Arial", 10, "bold")).pack(anchor="w", padx=40, pady=(10,5))
    end_cal = DateEntry(
        root,
        width=25,
        background="darkblue",
        foreground="white",
        borderwidth=2,
        date_pattern="dd.mm.yyyy",
        year=end_default.year,
        month=end_default.month,
        day=end_default.day,
    )
    end_cal.pack(padx=40, pady=5)

    tk.Label(root, text="Ответственный руководитель работ: Алимханов Идрис Махмудович",
             font=("Arial", 10, "bold"), fg="darkgreen").pack(anchor="w", padx=40, pady=15)

    btn = tk.Button(root, text="СОЗДАТЬ НАРЯД", command=generate,
                    bg="#0066cc", fg="white", font=("Arial", 14, "bold"), height=2, width=40)
    btn.pack(pady=30)

    # Метка имени файла
    filename_label = tk.Label(root, text="", fg="gray", font=("Arial", 9))
    filename_label.pack()

    # Обновление имени файла при изменении
    station_var.trace("w", update_filename_label)
    leader_combo.bind("<<ComboboxSelected>>", update_filename_label)

    update_filename_label()  # начальное обновление
    root.mainloop()


if __name__ == "__main__":
    run_gui()
