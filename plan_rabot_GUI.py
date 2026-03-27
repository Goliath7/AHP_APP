try:
    import tkinter as tk
    from tkinter import ttk, messagebox
    from tkcalendar import DateEntry
except Exception:  # pragma: no cover - for headless/web environments
    tk = None
    ttk = None
    messagebox = None
    DateEntry = None
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from shared_data import STATIONS as ALL_STATIONS, LEADERS_FULL as ALL_LEADERS, fio_short, output_path

# ====================== ДАННЫЕ ======================
STATIONS = ALL_STATIONS
SUPERVISORS = ALL_LEADERS

WORK_TEMPLATES = [
    {"name": "Прокладка силовых кабельных линий", "place": "Фасады станции", "volume": 443, "days": 9},
    {"name": "Установка щита электропитания и управления освещением", "place": "Щитовая", "volume": 2, "days": 8},
    {"name": "Монтаж несущих конструкций (сверление, установка анкеров или шпилек на химический анкер, установка металлических кронштейнов) и светотехнического оборудования", "place": "Фасады станции", "volume": 357, "days": 16},
    {"name": "Коммутация, юстировка и подключение светотехнического оборудования, подключение системы АХП к существующей системе электроснабжения метрополитена", "place": "Фасады станции и щитовая", "volume": 357, "days": 10},
    {"name": "Пуско-наладочные работы", "place": "Фасады станции и щитовая", "volume": "-", "days": 2},
    {"name": "Комплексное опробование светотехнического оборудования", "place": "Фасады станции и щитовая", "volume": "-", "days": 2},
]

MACHINES = [
    ['Аккумуляторная дрель', '3'], ['Аккумуляторный шуруповерт', '3'],
    ['Аккумуляторный перфоратор', '3'], ['Аккумуляторный клепатель', '2'],
    ['Аккумуляторная дрель-шуруповерт', '5']
]

def set_font(run, size=11, bold=False, italic=False, underline=False):
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.underline = underline
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Calibri')
    rFonts.set(qn('w:hAnsi'), 'Calibri')
    rPr.append(rFonts)

# ====================== ГЕНЕРАЦИЯ ======================
def generate_document(station=None, supervisor_full=None, start_dt=None, end_dt=None, show_messages=True):
    def fail(message):
        if show_messages and messagebox is not None:
            messagebox.showerror("Ошибка", message)
            return None
        raise ValueError(message)

    station = station_combo.get().strip() if station is None else station.strip()
    supervisor_full = supervisor_combo.get().strip() if supervisor_full is None else supervisor_full.strip()
    if not station or not supervisor_full:
        return fail("Выберите станцию и руководителя")

    start_dt = start_date_entry.get_date() if start_dt is None else start_dt
    end_dt = end_date_entry.get_date() if end_dt is None else end_dt
    if start_dt > end_dt:
        return fail("Дата начала позже окончания")

    total_days = (end_dt - start_dt).days + 1
    if total_days < 8:
        return fail("Период слишком короткий — минимум 8 дней для п.5+6+7")

    # Дни для графика
    days_list = []
    day_map = {}
    dt = start_dt
    idx = 0
    while dt <= end_dt:
        days_list.append(dt)
        day_map[dt] = idx
        dt += timedelta(days=1)
        idx += 1

    def day_label(index):
        day = days_list[index]
        return str(day.day)

    # ─── Планирование работ (номера 2–7, последние 4 дня — п.6+п.7) ───
    data1 = []

    # Фиксированные последние работы
    start_p7 = end_dt - timedelta(days=1)               # последние 2 дня
    end_p7   = end_dt
    start_p6 = start_p7 - timedelta(days=2)             # предпоследние 2 дня
    end_p6   = start_p7 - timedelta(days=1)

    end_p5   = start_p6 - timedelta(days=1)             # п.5 заканчивается за 4 дня до конца

    # п.7
    data1.append([
        "7", f'Ст. “{station}”', WORK_TEMPLATES[5]["name"], "", WORK_TEMPLATES[5]["place"], "7",
        str(WORK_TEMPLATES[5]["volume"]), f"{start_p7:%d.%m.%Y}\n{end_p7:%d.%m.%Y}"
    ])

    # п.6
    data1.append([
        "6", f'Ст. “{station}”', WORK_TEMPLATES[4]["name"], "", WORK_TEMPLATES[4]["place"], "7",
        str(WORK_TEMPLATES[4]["volume"]), f"{start_p6:%d.%m.%Y}\n{end_p6:%d.%m.%Y}"
    ])

    overlap = 3
    current_end = end_p5

    # п.5 — с перекрытием (как п.2–4)
    w = WORK_TEMPLATES[3]
    duration = w["days"]
    start_p5 = current_end - timedelta(days=duration - 1)
    start_p5 = max(start_p5, start_dt)
    data1.append([
        "5", f'Ст. “{station}”', w["name"], "", w["place"], "7",
        str(w["volume"]), f"{start_p5:%d.%m.%Y}\n{current_end:%d.%m.%Y}"
    ])
    current_end = start_p5 - timedelta(days=1 - overlap)
    if current_end < start_dt:
        current_end = start_dt

    # п.4, п.3, п.2 — с перекрытием
    for i in range(2, -1, -1):
        w = WORK_TEMPLATES[i]
        duration = w["days"]
        start_this = current_end - timedelta(days=duration - 1)
        start_this = max(start_this, start_dt)

        data1.append([
            str(i + 2), f'Ст. “{station}”', w["name"], "", w["place"], "7",
            str(w["volume"]), f"{start_this:%d.%m.%Y}\n{current_end:%d.%m.%Y}"
        ])

        current_end = start_this - timedelta(days=1 - overlap)
        if current_end < start_dt:
            current_end = start_dt

    data1.reverse()  # → порядок 2,3,4,5,6,7

    def parse_range(range_text):
        s, e = [x.strip() for x in range_text.split('\n')]
        return datetime.strptime(s, '%d.%m.%Y').date(), datetime.strptime(e, '%d.%m.%Y').date()

    first_start = parse_range(data1[0][7])[0]
    if first_start < start_dt and show_messages:
        messagebox.showwarning("Внимание", "Работы начинаются раньше указанной даты начала — объём не помещается.")

    # Если график начинается позже выбранной даты, сдвигаем его к началу периода.
    if first_start > start_dt:
        shift = first_start - start_dt
        for row in data1:
            row_start, row_end = parse_range(row[7])
            row_start -= shift
            row_end -= shift
            if row_start < start_dt:
                row_start = start_dt
            if row_end < row_start:
                row_end = row_start
            row[7] = f"{row_start:%d.%m.%Y}\n{row_end:%d.%m.%Y}"

    # ─── Документ ───
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(1.5)
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)

    # РАЗРАБОТАЛ / УТВЕРЖДАЮ
    sig_table = doc.add_table(1, 2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig_table.columns[0].width = Cm(9.5)
    sig_table.columns[1].width = Cm(9.5)

    for i, title in enumerate(["РАЗРАБОТАЛ:", "УТВЕРЖДАЮ:"]):
        p = sig_table.cell(0, i).paragraphs[0]
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run(title)
        set_font(run, 11, bold=True)
        p.add_run('\nНачальник участка\nООО «Азбука Света»\n_______________________ / Алимханов И. М. /\n' if i == 0 else
                  '\nГенеральный директор\nООО «Азбука Света»\n_______________________ / Макаров А. А. /\n')
        p.add_run(f'__________{datetime.now().year}г.')

    for _ in range(5): doc.add_paragraph()

    # ПЛАН РАБОТ
    p = doc.add_paragraph("ПЛАН РАБОТ")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_font(p.runs[0], 28, bold=True)

    p = doc.add_paragraph()
    p.add_run("Объект:           ")
    run = p.add_run(f'Станция «{station}» Московского метрополитена')
    set_font(run, 12, underline=True)
    doc.add_paragraph("(наименование)").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    p = doc.add_paragraph()
    p.add_run("Период выполнения работ          ")
    run = p.add_run(f"с {start_dt:%d.%m.%Y}г. по {end_dt:%d.%m.%Y}г.")
    set_font(run, 12, underline=True)

    for text in ["Заказчик", "Сторонняя организация"]:
        p = doc.add_paragraph(text + "          ")
        run = p.add_run("Служба Электроснабжения ДО-3 ГУП «Московский метрополитен»" if text == "Заказчик" else "ООО «Азбука Света»")
        set_font(run, 12, underline=True)
        doc.add_paragraph("(наименование)").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for _ in range(7): doc.add_paragraph()

    # СОГЛАСОВАНО
    sig_table = doc.add_table(1, 2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for col in (0, 1):
        p = sig_table.cell(0, col).paragraphs[0]
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = p.add_run('СОГЛАСОВАНО')
        set_font(run, 12, bold=True)
        p.add_run('\nРуководитель причастного подразделения\nГУП «Московский метрополитен»\n__________________________ Ф.И.О\n«____» _______________ 2026 г.')

    doc.add_paragraph()

    # Таблица основных работ (теперь с номерами 2–7 и правильными периодами)
    p = doc.add_paragraph(f'по строительно-монтажным работам на ст. «{station}»')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_font(p.runs[0], 12)

    headers1 = ['№ п/п', 'Станция', 'Наименование работ', 'Количество помещений',
                'Место проведения работ', 'Кол-во рабочего персонала', 'Объем работ', 'Время выполнения']

    tbl1 = doc.add_table(rows=1 + len(data1), cols=len(headers1))
    tbl1.style = 'Table Grid'
    tbl1.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers1):
        cell = tbl1.rows[0].cells[i]
        cell.text = h
        set_font(cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(h), 9, bold=True)

    for r, row_data in enumerate(data1, 1):
        for c, value in enumerate(row_data):
            cell = tbl1.rows[r].cells[c]
            cell.text = str(value)
            set_font(cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(str(value)), 9)

    doc.add_paragraph()

    # Примечание, описание работ, порядок и т.д. — оставляем как было
    p = doc.add_paragraph('Примечание: ')
    set_font(p.runs[0], 12, bold=True)
    p.add_run('\nРаботы выполняются круглосуточно с предварительного согласования со смежными службами при условии отсутствия помех для движения пассажиров\n'
              'Доставка и перемещение необходимых материалов и инструментов осуществляется вручную ежесменно в объеме сменной потребности\n'
              'Количество рабочего персонала рассчитано на установленный объем, согласно графику производства работ.\n'
              'Движение основных строительных машин, а также выделение хозяйственных мотоединиц не требуется.   \n'
              'Основание для проведения работ по договору № 4313м от 10.06.2024г. (до 25.12.2026г. включительно) и Приказ № УД-07-2127/24/17 от 24.06.2024г. (до 10.06.2026г. включительно)')
    set_font(p.runs[1], 11)
    doc.add_paragraph()
    p = doc.add_paragraph(),
    p = doc.add_paragraph('Описание работ')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    set_font(p.runs[0], 12, bold=True)
    for item in [
        '2) Прокладка силовых кабельных линий;',
        '3) Установка щита электропитания и управления освещением;',
        '4) Монтаж несущих конструкций (сверление, установка анкеров или шпилек на химический анкер, установка металлических кронштейнов) и светотехнического оборудования;',
        '5) Коммутация, юстировка и подключение светотехнического оборудования; подключение системы АХП к существующей системе электроснабжения метрополитена',
        '6) Пуско-наладочные работы;',
        '7) Комплексное опробование светотехнического оборудования.'
    ]:
        doc.add_paragraph(item)
        set_font(doc.paragraphs[-1].runs[0], 12)


    doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    # ====================== ПОРЯДОК ДОСТАВКИ ======================
    p = doc.add_paragraph('Порядок доставки на объект строительных конструкций, изделий, материалов и оборудования')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    set_font(p.runs[0], 12, bold=True)
    p = doc.add_paragraph('Строительные инструменты, оборудование и приспособления, а также материал доставляется на объект с помощью автомашин, разгрузка производится вручную, либо с помощью использования рохли силами рабочих ООО «Азбука Света» и складируется в специально отведенных местах, с соблюдением требований «Правил по охране труда при погрузочно-разгрузочных работах и размещению грузов», утверждённых Приказом Минтруда РФ от 28.10.2020г. №75ВН. Точка подключения электрооборудования не требуется.')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    set_font(p.runs[0], 12)

    # ====================== ПОРЯДОК ПОДКЛЮЧЕНИЯ ======================
    p = doc.add_paragraph('Порядок подключения электроагрегатов и электропотребителей на объекте')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    set_font(p.runs[0], 12, bold=True)
    p = doc.add_paragraph('Для обеспечения необходимого освещения в местах производства работ устанавливаются переносные источники света марки ПСМ-50-1 (аккумуляторные светодиодные светильники)\n'
            'Устройство и эксплуатация электроустановок должны осуществляться в соответствии с требованиями правил устройства электроустановок, правил техники безопасности при эксплуатации электроустановок потребителей.')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    set_font(p.runs[0], 12)

    # ====================== ХРАНЕНИЕ ======================
    p = doc.add_paragraph('Хранение строительных материалов')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    set_font(p.runs[0], 12, bold=True)
    p = doc.add_paragraph('Хранение на объекте строительных материалов, инструментов и приспособлений, согласно ППР-АС-АХП-28-06.24. Ответственный руководитель работ в конце смены проводит детальный осмотр места производства работ на предмет обеспечения порядка.')     
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    set_font(p.runs[0], 12)
    p = doc.add_paragraph('Хранение на объекте бытового мусора, упаковок, тары и любых иных емкостей категорически запрещено.\n'
            'Проведение огневых работ не предусмотрено ')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    set_font(p.runs[0], 12)
    doc.add_paragraph()
    doc.add_paragraph()


    # График производства работ (альбомный, с перекрытиями)
    landscape = doc.add_section()
    landscape.orientation = WD_ORIENT.LANDSCAPE
    landscape.page_width = Cm(29.7)
    landscape.page_height = Cm(21)

    p = doc.add_paragraph('График производства работ по дням')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_font(p.runs[0], 14, bold=True)

    left_cols = 4
    tbl = doc.add_table(rows=2 + len(data1), cols=left_cols + total_days)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    tbl.columns[0].width = Cm(1.4)
    tbl.columns[1].width = Cm(13.0)
    tbl.columns[2].width = Cm(2.4)
    tbl.columns[3].width = Cm(2.4)
    for i in range(left_cols, left_cols + total_days):
        tbl.columns[i].width = Cm(0.75)

    row0 = tbl.rows[0].cells
    row0[0].text = '№\nп/п'
    row0[1].text = 'Наименование'
    row0[2].text = 'Ед.\nизм.'
    row0[3].text = 'Кол-\nво'
    for idx in range(4):
        p = row0[idx].paragraphs[0]
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        if p.runs:
            set_font(p.runs[0], 8, bold=True)
    merged = row0[left_cols].merge(row0[-1])
    p_days = merged.paragraphs[0]
    p_days.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_font(p_days.add_run('График производства работ по дням'), 8, bold=True)

    row1 = tbl.rows[1].cells
    row1[1].text = 'Перечень\nпроизводимых\nработ'
    p = row1[1].paragraphs[0]
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_font(p.runs[0], 8, bold=True)

    for idx, _ in enumerate(days_list):
        col_idx = left_cols + idx
        row1[col_idx].text = day_label(idx)
        set_font(row1[col_idx].paragraphs[0].runs[0], 8, bold=True)

    gray = "D9D9D9"
    for r_idx, row_data in enumerate(data1, 2):
        row = tbl.rows[r_idx]
        for c, value in enumerate([row_data[0], row_data[2], "", row_data[6]]):
            cell = row.cells[c]
            cell.text = str(value)
            p = cell.paragraphs[0]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER if c != 1 else WD_PARAGRAPH_ALIGNMENT.LEFT
            set_font(p.runs[0], 8)

        try:
            s, e = [x.strip() for x in row_data[7].split('\n')]
            start = datetime.strptime(s, '%d.%m.%Y').date()
            end = datetime.strptime(e, '%d.%m.%Y').date()
            cur = start
            while cur <= end:
                if cur in day_map:
                    col = left_cols + day_map[cur]
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:fill'), gray)
                    tcPr = row.cells[col]._tc.get_or_add_tcPr()
                    tcPr.append(shading)
                cur += timedelta(days=1)
        except:
            pass
    p = doc.add_paragraph()
    p = doc.add_paragraph(),
    # ====================== ГРАФИК ДВИЖЕНИЯ ТРУДОВЫХ РЕСУРСОВ ======================
    # (используем оригинальную структуру из plan_rabot.py с динамическими днями)
    p = doc.add_paragraph('График движения трудовых ресурсов')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_font(p.runs[0], 14, bold=True)

    left_cols = 3
    total_cols = left_cols + total_days

    tbl3 = doc.add_table(rows=3 + 2, cols=total_cols)
    tbl3.style = 'Table Grid'
    tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl3.allow_autofit = False

    tbl3.columns[0].width = Cm(1.4)
    tbl3.columns[1].width = Cm(13.0)
    tbl3.columns[2].width = Cm(3.0)
    for i in range(left_cols, total_cols):
        tbl3.columns[i].width = Cm(0.75)

    days_header_cell = tbl3.cell(0, left_cols)
    days_header_cell.merge(tbl3.cell(0, total_cols - 1))
    p_days = days_header_cell.paragraphs[0]
    p_days.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_days = p_days.add_run('Средняя численность рабочих по дням')
    set_font(run_days, 10)

    hdr_row = tbl3.rows[1].cells
    hdr_row[0].text = '№'
    hdr_row[1].text = 'Наименование\nпрофессий рабочих'
    hdr_row[2].text = 'Численность\nрабочих'
    for idx, _ in enumerate(days_list):
        hdr_row[left_cols + idx].text = day_label(idx)

    for cell in hdr_row:
        p = cell.paragraphs[0]
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        set_font(p.runs[0] if p.runs else p.add_run(cell.text), 10)

    for col in range(left_cols):
        tbl3.cell(0, col).merge(tbl3.cell(1, col))

    tbl3.cell(1, 0).text = '№'
    tbl3.cell(1, 1).text = 'Наименование\nпрофессий рабочих'
    tbl3.cell(1, 2).text = 'Численность\nрабочих'

    for col in range(left_cols):
        p = tbl3.cell(1, col).paragraphs[0]
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        set_font(p.runs[0] if p.runs else p.add_run(p.text), 11)

    labor_data = [
        ['1', 'ИТР', '2', [2] * total_days],
        ['2', 'Монтажники', '5', [5] * total_days]
    ]

    for r, ld in enumerate(labor_data):
        row = tbl3.rows[2 + r]
        row.cells[0].text = ld[0]
        row.cells[1].text = ld[1]
        row.cells[2].text = ld[2]
        for c in range(left_cols):
            p = row.cells[c].paragraphs[0]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            set_font(p.runs[0] if p.runs else p.add_run(row.cells[c].text), 10)
        for c_idx in range(left_cols, total_cols):
            val = str(ld[3][c_idx - left_cols])
            row.cells[c_idx].text = val
            p = row.cells[c_idx].paragraphs[0]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            set_font(p.runs[0] if p.runs else p.add_run(val), 9)

    doc.add_paragraph()

    # ====================== ГРАФИК ПОТРЕБНОСТИ В МАШИНАХ ======================
    p = doc.add_paragraph('График потребности в строительных машинах и электроинструменте')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_font(p.runs[0], 14, bold=True)

    left_cols_m = 2
    total_cols_m = left_cols_m + total_days

    tbl4 = doc.add_table(rows=2 + len(MACHINES), cols=total_cols_m)
    tbl4.style = 'Table Grid'
    tbl4.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl4.allow_autofit = False

    tbl4.columns[0].width = Cm(14.0)
    tbl4.columns[1].width = Cm(3.5)
    for i in range(left_cols_m, total_cols_m):
        tbl4.columns[i].width = Cm(0.75)

    merged_days = tbl4.cell(0, left_cols_m)
    merged_days.merge(tbl4.cell(0, total_cols_m - 1))
    p_days = merged_days.paragraphs[0]
    p_days.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_days = p_days.add_run('Среднесуточная численность машин по дням')
    set_font(run_days, 10, bold=True)

    hdr = tbl4.rows[1].cells
    hdr[0].text = 'Наименование'
    hdr[1].text = 'Кол-во машин'
    for idx, _ in enumerate(days_list):
        hdr[left_cols_m + idx].text = day_label(idx)

    for cell in hdr:
        p = cell.paragraphs[0]
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        set_font(p.runs[0] if p.runs else p.add_run(cell.text), 10)

    for col in range(left_cols_m):
        tbl4.cell(0, col).merge(tbl4.cell(1, col))

    tbl4.cell(1, 0).text = 'Наименование'
    tbl4.cell(1, 1).text = 'Кол-во машин'

    for col in range(left_cols_m):
        p = tbl4.cell(1, col).paragraphs[0]
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        set_font(p.runs[0] if p.runs else p.add_run(p.text), 11)

    for r, (name, qty) in enumerate(MACHINES, start=2):
        row = tbl4.rows[r]
        row.cells[0].text = name
        row.cells[1].text = qty
        for c in range(left_cols_m):
            p = row.cells[c].paragraphs[0]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT if c == 0 else WD_PARAGRAPH_ALIGNMENT.CENTER
            set_font(p.runs[0] if p.runs else p.add_run(row.cells[c].text), 10)
        for c_idx in range(left_cols_m, total_cols_m):
            cell = row.cells[c_idx]
            cell.text = qty
            p = cell.paragraphs[0]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            set_font(p.runs[0] if p.runs else p.add_run(qty), 9)

    doc.add_paragraph()


    # Руководитель
    short_sup = fio_short(supervisor_full)
    p = doc.add_paragraph(f"Руководитель работ ООО «Азбука Света» ______________________________________ \\ {short_sup} \\")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    set_font(p.runs[0], 12)

    # Лист ознакомления
    portrait_sec = doc.add_section()
    portrait_sec.orientation = WD_ORIENT.PORTRAIT
    portrait_sec.page_width = Cm(21.0)
    portrait_sec.page_height = Cm(29.7)
    portrait_sec.left_margin = Cm(2.5)
    portrait_sec.right_margin = Cm(2.0)
    portrait_sec.top_margin = Cm(2.0)
    portrait_sec.bottom_margin = Cm(2.0)

    p = doc.add_paragraph('Лист ознакомления с планом работ')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_font(p.runs[0], 14, bold=True)

    doc.add_paragraph()

    today_str = datetime.now().strftime('%d.%m.%Y')
    fam_data = [
        ['1', 'Макаров Артём Андреевич', 'Генеральный директор', '', today_str],
        ['2', 'Алимханов Идрис Махмудович', 'Ответственный руководитель работ', '', today_str],
        ['3', supervisor_full, 'Руководитель работ', '', today_str],
        ['4', 'Верзилин Андрей Игоревич', 'Электромонтажник', '', today_str],
        ['5', 'Агроматин Игорь Сергеевич', 'Электромонтажник', '', today_str],
        ['6', 'Осин Максим Васильевич', 'Электромонтажник', '', today_str],
        ['7', 'Баязов Асланбек Маирбекович', 'Электромонтажник', '', today_str],
    ]

    fam_headers = ['№ п/п', 'Ф.И.О', 'Должность', 'Подпись', 'Дата']
    tbl_fam = doc.add_table(rows=1 + len(fam_data), cols=len(fam_headers))
    tbl_fam.style = 'Table Grid'
    tbl_fam.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_fam.allow_autofit = False

    tbl_fam.columns[0].width = Cm(2.0)
    tbl_fam.columns[1].width = Cm(8.0)
    tbl_fam.columns[2].width = Cm(6.0)
    tbl_fam.columns[3].width = Cm(6.0)
    tbl_fam.columns[4].width = Cm(4.0)

    hdr = tbl_fam.rows[0].cells
    for i, text in enumerate(fam_headers):
        hdr[i].text = text
        set_font(hdr[i].paragraphs[0].runs[0] if hdr[i].paragraphs[0].runs else hdr[i].paragraphs[0].add_run(text), 11, bold=True)

    for r_idx, row_d in enumerate(fam_data):
        row = tbl_fam.rows[r_idx + 1]
        for c_idx, text in enumerate(row_d):
            cell = row.cells[c_idx]
            cell.text = text
            p_cell = cell.paragraphs[0]
            p_cell.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER if c_idx in (0, 4) else WD_PARAGRAPH_ALIGNMENT.LEFT
            set_font(p_cell.runs[0] if p_cell.runs else p_cell.add_run(text), 11)
        row.height = Cm(1.6)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    filename = f"План_работ_{station.replace(' ', '_')}_{short_sup}_{start_dt.strftime('%d.%m')}-{end_dt.strftime('%d.%m')}.docx"
    file_path = output_path(filename)
    doc.save(file_path)
    if show_messages:
        messagebox.showinfo("Готово", f"Файл сохранён:\n{file_path}")
    return str(file_path)

# ====================== GUI ======================
def run_gui():
    global root, station_combo, supervisor_combo, start_date_entry, end_date_entry
    if tk is None or ttk is None or DateEntry is None:
        raise RuntimeError("Tkinter GUI недоступен в этой среде")

    root = tk.Tk()
    root.title("Генератор плана работ")
    root.geometry("560x340")

    tk.Label(root, text="Станция:").grid(row=0, column=0, padx=10, pady=8, sticky="e")
    station_combo = ttk.Combobox(root, values=STATIONS, width=40, state="readonly")
    station_combo.grid(row=0, column=1, pady=8)

    tk.Label(root, text="Руководитель:").grid(row=1, column=0, padx=10, pady=8, sticky="e")
    supervisor_combo = ttk.Combobox(root, values=SUPERVISORS, width=40, state="readonly")
    supervisor_combo.grid(row=1, column=1, pady=8)

    tk.Label(root, text="Начало:").grid(row=2, column=0, padx=10, pady=8, sticky="e")
    start_date_entry = DateEntry(root, width=14, date_pattern='dd.mm.yyyy')
    start_date_entry.grid(row=2, column=1, pady=8, sticky="w")

    tk.Label(root, text="Окончание:").grid(row=3, column=0, padx=10, pady=8, sticky="e")
    end_date_entry = DateEntry(root, width=14, date_pattern='dd.mm.yyyy')
    end_date_entry.grid(row=3, column=1, pady=8, sticky="w")

    ttk.Button(root, text="Создать документ", command=generate_document).grid(row=5, column=0, columnspan=2, pady=20)
    root.mainloop()


if __name__ == "__main__":
    run_gui()
